#!/usr/bin/env python3
"""
Reference script: export simple Markdown to DOCX.
Status: reference only, not production-ready.
Requires: python-docx
Usage: python export_docx_from_markdown.py input.md output.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError as exc:
    raise SystemExit("Missing dependency: python-docx. Install with `pip install python-docx`.") from exc


def add_paragraph_with_style(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        doc.add_paragraph("")
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
    elif stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
    elif stripped.startswith("- [ ] "):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ " + stripped[6:].strip())
    elif stripped.startswith("- [x] ") or stripped.startswith("- [X] "):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☑ " + stripped[6:].strip())
    elif stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
    elif stripped.startswith("> "):
        p = doc.add_paragraph()
        run = p.add_run(stripped[2:].strip())
        run.italic = True
    else:
        doc.add_paragraph(stripped)


def export_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    if output_path.exists():
        raise FileExistsError(f"Output file already exists. Refusing to overwrite: {output_path}")

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    content = input_path.read_text(encoding="utf-8")
    in_code_block = False
    code_lines: list[str] = []

    for line in content.splitlines():
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        if in_code_block:
            code_lines.append(line)
        else:
            add_paragraph_with_style(doc, line)

    doc.save(output_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: python export_docx_from_markdown.py input.md output.docx")
    export_markdown_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
