import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_preset_style(doc, preset_name):
    # This is a programmatic styling approach
    # We define minimal styling dynamically instead of using heavy DOCX templates
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    if preset_name == 'vn-gov':
        font.name = 'Times New Roman'
        font.size = Pt(14)
    elif preset_name == 'mckinsey':
        font.name = 'Arial'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    elif preset_name == 'modern-minimal':
        font.name = 'Helvetica'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    # Set heading styles
    for i in range(1, 4):
        heading_style = doc.styles.get(f'Heading {i}')
        if heading_style:
            heading_font = heading_style.font
            if preset_name == 'vn-gov':
                heading_font.name = 'Times New Roman'
                heading_font.color.rgb = RGBColor(0, 0, 0)
            else:
                heading_font.color.rgb = RGBColor(0x00, 0x33, 0x66) # Default dark blue

def convert_md_to_docx(md_path, docx_path, preset_name='classic-corporate'):
    try:
        # Create a blank document (0 bytes required as base)
        doc = Document()
        apply_preset_style(doc, preset_name)
        
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
                
        doc.save(docx_path)
        print(f"Generated Lightweight DOCX saved to {docx_path} using preset '{preset_name}'")
        return True
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx_engine.py <input.md> <output.docx> [preset_name]")
        sys.exit(1)
        
    input_md = sys.argv[1]
    output_docx = sys.argv[2]
    preset = sys.argv[3] if len(sys.argv) > 3 else "classic-corporate"
    
    convert_md_to_docx(input_md, output_docx, preset)
